from flask import (
    Flask,
    request,
    render_template,
    jsonify,
    session,
    redirect,
    url_for,
    send_file
)

from flask_cors import CORS
from docx import Document
from datetime import datetime
from werkzeug.utils import secure_filename
from authlib.integrations.flask_client import OAuth

from groq import Groq

import os
import tempfile
import subprocess
import platform
import urllib.parse

# =====================================================
# APP CONFIG
# =====================================================

app = Flask(__name__, template_folder="templates")
CORS(app)

app.secret_key = os.environ.get("SECRET_KEY", "supersecret")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# =====================================================
# GROQ AI
# =====================================================

groq_client = Groq(
    api_key=os.environ.get("Alfa_Letters")
)

# =====================================================
# GOOGLE OAUTH
# =====================================================

oauth = OAuth(app)

google = oauth.register(
    name="google",
    client_id=os.environ.get("GOOGLE_CLIENT_ID"),
    client_secret=os.environ.get("GOOGLE_CLIENT_SECRET"),
    server_metadata_url="https://accounts.google.com/.well-known/openid-configuration",
    client_kwargs={"scope": "openid email profile"}
)

ALLOWED_EMAILS = ["hr@alfatza.com"]

# =====================================================
# TEMPLATE FILES
# =====================================================

TEMPLATE_FILES = {

    # OFFER LETTERS
    "offer_telecaller":
        os.path.join(BASE_DIR, "templates_docx", "offer_telecaller.docx"),

    "offer_team_leader":
        os.path.join(BASE_DIR, "templates_docx", "offer_team_leader.docx"),

    "offer_backend":
        os.path.join(BASE_DIR, "templates_docx", "offer_backend.docx"),

    "offer_hr":
        os.path.join(BASE_DIR, "templates_docx", "offer_hr.docx"),

    "offer_data_analyst":
        os.path.join(BASE_DIR, "templates_docx", "offer_data_analyst.docx"),

    # INCREMENT
    "increment":
        os.path.join(BASE_DIR, "templates_docx", "increment.docx"),

    # EXPERIENCE
    "experience":
        os.path.join(BASE_DIR, "templates_docx", "experience.docx"),

    # TERMINATION
    "termination":
        os.path.join(BASE_DIR, "templates_docx", "termination.docx"),

    # ABSCOND
    "abscond":
        os.path.join(BASE_DIR, "templates_docx", "abscond.docx"),
}

# =====================================================
# BRANCHES
# =====================================================

BRANCHES = {
    "vashi": "Vashi Branch Address",
    "thane": "Thane Branch Address",
    "virar": "Virar Branch Address"
}

# =====================================================
# HELPERS
# =====================================================

def format_date(date_str):
    return datetime.strptime(
        date_str,
        "%Y-%m-%d"
    ).strftime("%d %B %Y")


def replace_text(doc, values):

    for para in doc.paragraphs:
        for k, v in values.items():
            if k in para.text:
                para.text = para.text.replace(k, v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in values.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)


def convert_to_pdf(docx_path, output_dir):

    libreoffice = "soffice"

    if platform.system() == "Windows":
        libreoffice = r"C:\Program Files\LibreOffice\program\soffice.exe"

    subprocess.run([
        libreoffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        output_dir,
        docx_path
    ], check=True)

    return os.path.join(
        output_dir,
        os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
    )


def build_gmail_link(to, subject, body):

    base = "https://mail.google.com/mail/?view=cm&fs=1"

    params = {
        "to": to,
        "su": subject,
        "body": body
    }

    return base + "&" + urllib.parse.urlencode(params)


def generate_ai_content(letter_type, data):

    prompt = f"""
You are an HR legal letter writer.

Write a professional {letter_type} letter paragraph.

EMPLOYEE NAME:
{data['name']}

ROLE:
{data['role']}

BRANCH:
{data['branch']}

RAW NOTES:
{data['ai_prompt']}

Rules:
- Professional HR tone
- Formal English
- 250-400 words
- Do not use placeholders
- Ready to paste directly into official company document
"""

    completion = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.7
    )

    return completion.choices[0].message.content


def get_template_path(letter_type, role):

    if letter_type == "offer":
        return TEMPLATE_FILES.get(f"offer_{role}")

    return TEMPLATE_FILES.get(letter_type)


def get_download_name(name, letter_type):

    clean_type = letter_type.capitalize()

    return f"{name}_{clean_type}_Letter.pdf"


# =====================================================
# AUTH ROUTES
# =====================================================

@app.route("/login")
def login():
    return google.authorize_redirect(
        url_for("authorize", _external=True)
    )


@app.route("/authorize")
def authorize():

    token = google.authorize_access_token()

    user = token.get("userinfo")

    if not user:
        return "Login failed"

    email = user.get("email")

    if email not in ALLOWED_EMAILS:
        return "Unauthorized"

    session["user"] = user

    return redirect("/")


@app.route("/logout")
def logout():

    session.clear()

    return redirect("/")


# =====================================================
# HOME
# =====================================================

@app.route("/")
def home():

    if "user" not in session:
        return redirect("/login")

    return render_template(
        "index.html",
        user=session["user"]
    )


# =====================================================
# MAIN GENERATE ROUTE
# =====================================================

@app.route("/generate", methods=["POST"])
def generate():

    try:

        if "user" not in session:
            return jsonify({
                "error": "Unauthorized"
            }), 401

        data = request.get_json()

        # =================================================
        # BASIC REQUIRED
        # =================================================

        required = [
            "name",
            "employee_code",
            "phone",
            "email",
            "address",
            "role",
            "branch",
            "joining",
            "letter_type"
        ]

        for field in required:

            if not data.get(field):
                return jsonify({
                    "error": f"Missing {field}"
                }), 400

        # =================================================
        # OFFER / INCREMENT NEED SALARY
        # =================================================

        if data["letter_type"] in ["offer", "increment"]:

            if not data.get("salary"):
                return jsonify({
                    "error": "Salary required"
                }), 400

        # =================================================
        # TEMPLATE
        # =================================================

        template_path = get_template_path(
            data["letter_type"],
            data["role"]
        )

        if not template_path or not os.path.exists(template_path):

            return jsonify({
                "error": "Template not found"
            }), 500

        # =================================================
        # LOAD DOC
        # =================================================

        doc = Document(template_path)

        # =================================================
        # AI GENERATED CONTENT
        # =================================================

        ai_content = ""

        if data["letter_type"] in [
            "termination",
            "abscond",
            "experience"
        ]:

            ai_content = generate_ai_content(
                data["letter_type"],
                data
            )

        # =================================================
        # VALUES
        # =================================================

        values = {

            "{{name}}":
                data["name"],

            "{{employee_code}}":
                data["employee_code"],

            "{{phone}}":
                data["phone"],

            "{{email}}":
                data["email"],

            "{{address}}":
                data["address"],

            "{{role}}":
                data["role"].replace("_", " ").title(),

            "{{branch_address}}":
                BRANCHES.get(data["branch"], ""),

            "{{salary}}":
                data.get("salary", ""),

            "{{increment_salary}}":
                data.get("increment_salary", ""),

            "{{joining}}":
                format_date(data["joining"]),

            "{{date}}":
                datetime.now().strftime("%d %B %Y"),

            "{{ai_content}}":
                ai_content
        }

        replace_text(doc, values)

        # =================================================
        # SAVE FILE
        # =================================================

        temp_dir = tempfile.mkdtemp()

        safe_name = secure_filename(data["name"])

        file_base = f"{safe_name}_{data['letter_type']}"

        docx_path = os.path.join(
            temp_dir,
            f"{file_base}.docx"
        )

        doc.save(docx_path)

        # =================================================
        # CONVERT PDF
        # =================================================

        pdf_path = convert_to_pdf(
            docx_path,
            temp_dir
        )

        # =================================================
        # MAIL SUBJECT + BODY
        # =================================================

        branch_name = data["branch"].capitalize()

        letter_name = data["letter_type"].capitalize()

        subject = (
            f"{letter_name} Letter "
            f"- {branch_name} Branch"
        )

        body = f"""Dear {data['name']},

Please find attached your formal {letter_name} Letter for your records.

Kindly review the document carefully and acknowledge receipt.

For any clarification, feel free to contact the HR department.

Warm regards,
Rashid Ali
H.R
ALFA TZA LLP
"""

        gmail_link = build_gmail_link(
            data["email"],
            subject,
            body
        )

        # =================================================
        # STORE DOWNLOAD
        # =================================================

        session[file_base] = pdf_path

        # =================================================
        # RESPONSE
        # =================================================

        return jsonify({
            "success": True,

            "download_url":
                f"/download/{file_base}",

            "gmail_link":
                gmail_link
        })

    except subprocess.CalledProcessError:

        return jsonify({
            "error": "PDF conversion failed"
        }), 500

    except Exception as e:

        return jsonify({
            "error": str(e)
        }), 500


# =====================================================
# DOWNLOAD
# =====================================================

@app.route("/download/<filename>")
def download(filename):

    try:

        pdf_path = session.get(filename)

        if not pdf_path:
            return "File not found", 404

        clean_name = filename.replace("_", " ")

        return send_file(
            pdf_path,
            as_attachment=True,
            download_name=f"{clean_name}.pdf"
        )

    except Exception:
        return "Download failed", 500


# =====================================================
# RUN
# =====================================================

if __name__ == "__main__":
    app.run(debug=True)
